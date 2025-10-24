import pandas as pd
from docx import Document
from docx.shared import Pt
from typing import Dict


def calculate_fault_type_proportions(csv_path: str, output_docx: str) -> None:
    """
    读取CSV文件，统计每个file_name对应的fault_type占比，
    """
    # 1. 读取CSV文件
    df = pd.read_csv(csv_path)

    if "file_name" not in df.columns or "fault_type" not in df.columns:
        raise ValueError("输入的CSV文件必须包含 'file_name' 和 'fault_type' 列。")

    # 2. 统计每个file_name下各fault_type的数量并计算占比
    counts = df.groupby(["file_name", "fault_type"]).size().reset_index(name="count")
    total_counts = counts.groupby("file_name")["count"].transform("sum")
    counts["proportion"] = counts["count"] / total_counts

    # 转换为透视表，方便输出为表格
    pivot_df = counts.pivot_table(
        index="file_name",
        columns="fault_type",
        values="proportion",
        fill_value=0
    )

    # 计算每个file_name的最可能故障类型（占比最大）
    most_likely = pivot_df.idxmax(axis=1)

    # 3. 写入Word文档
    doc = Document()
    doc.add_heading("每个文件对应各类故障类型占比统计表", level=1)

    # 创建表格
    table = doc.add_table(rows=1, cols=len(pivot_df.columns) + 2, style="Table Grid")

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "file_name"
    for i, col in enumerate(pivot_df.columns, start=1):
        hdr_cells[i].text = col
    hdr_cells[len(pivot_df.columns) + 1].text = "最可能故障类型"

    # 填充数据
    for file_name, row in pivot_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(file_name)
        for i, col in enumerate(pivot_df.columns, start=1):
            row_cells[i].text = f"{row[col]:.2%}"  # 百分比格式
        row_cells[len(pivot_df.columns) + 1].text = most_likely[file_name]

    # 美化字体大小
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # 保存Word文件
    doc.save(output_docx)
    print(f"结果已成功保存到: {output_docx}")

if __name__ == "__main__":
    input_csv = "target_with_predictions.csv"  # 输入CSV路径
    output_word = "fault_type_proportions.docx"       # 输出Word路径
    calculate_fault_type_proportions(input_csv, output_word)
