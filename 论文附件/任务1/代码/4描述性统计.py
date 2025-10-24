import pandas as pd
from docx import Document

# 1. 读取CSV文件
file_path = "source_domain_features_resampled_32k.csv"
df = pd.read_csv(file_path)

# 2. 仅保留数值型数据并计算描述性统计
desc_stats: pd.DataFrame = df.describe().T  # 只针对数值型变量
desc_stats.reset_index(inplace=True)
desc_stats.rename(columns={"index": "变量名"}, inplace=True)

# 3. 创建Word文档
doc = Document()
doc.add_heading("数值型变量的描述性统计结果", level=1)

# 添加表格
rows, cols = desc_stats.shape
table = doc.add_table(rows=rows + 1, cols=cols)
table.style = "Table Grid"

# 表头
for j, col in enumerate(desc_stats.columns):
    table.cell(0, j).text = str(col)

# 填充数据
for i in range(rows):
    for j in range(cols):
        value = desc_stats.iloc[i, j]
        if pd.isna(value):
            text = ""
        elif isinstance(value, float):
            text = f"{value:.4f}"
        else:
            text = str(value)
        table.cell(i + 1, j).text = text

# 保存Word文档
output_path = "描述性统计.docx"
doc.save(output_path)

print(f"文档已生成：{output_path}")
